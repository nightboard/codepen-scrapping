from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.proxy import Proxy, ProxyType

from docx import Document

from codepen import get_codepen_data

DATA_FILE = "links.txt"

options = webdriver.FirefoxOptions()
options.headless = True
driver = webdriver.Firefox(options=options, executable_path="./geckodriver")

# set loding time for website & send request
driver.implicitly_wait(10)

# document configuration
document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Sans Serif'
style.paragraph_format.line_spacing = 0.5


def add_meta_data(document, property_name, property_value=""):
    paragraph = document.add_paragraph('')
    run = paragraph.add_run(property_name)
    run.bold = True
    paragraph.add_run(property_value)


i = 1
with open(DATA_FILE) as read_file:
    for url in read_file:
        URL = url.replace('\n', '')

        codepen_data = get_codepen_data(driver, URL)

        add_meta_data(document, f"{i}. Title: ", codepen_data['title'])
        add_meta_data(document, "Author: ", codepen_data['author-name'])
        add_meta_data(document, "Created On: ", codepen_data['created-date'])
        add_meta_data(document, "Made with: ", codepen_data['made-with'])
        add_meta_data(document, "Responsive: ")
        add_meta_data(document, "Compatible browsers: ")
        add_meta_data(document, "Description: ")
        add_meta_data(document, "Code source link: ", f"{URL}\n\n\n")

        print(f"{i} ✅")
        i += 1
    document.save("demo.docx")

driver.close()
